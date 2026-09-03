#!/usr/bin/env python3
"""BukSU CMS-V2: Capstone Completion Certificate Template Generator
Generates a formal, ornamental, high-resolution .docx certificate template
triggered upon Phase 6 ADM verification.
"""

from pathlib import Path
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

WORKSPACE_ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DOCX = WORKSPACE_ROOT / "docs" / "templates" / "capstone-completion-certificate-template.docx"


def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)


def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>'
    )
    tcPr.append(tcMar)


def build_completion_certificate():
    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc = docx.Document()

    # Landscape or Standard Letter with tight margins for ornamental framing
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    # Base Styles
    doc.styles['Normal'].font.name = 'Georgia'
    doc.styles['Normal'].font.size = Pt(11)
    doc.styles['Normal'].font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

    # Ornamental Outer Table Frame
    outer_table = doc.add_table(rows=1, cols=1)
    outer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    outer_cell = outer_table.rows[0].cells[0]
    outer_cell.width = Inches(7.1)
    set_cell_background(outer_cell, "FDFBF7")  # Warm certificate parchment tint
    set_cell_margins(outer_cell, top=200, bottom=200, left=240, right=240)

    # Add XML double-line border to cell
    tcPr = outer_cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="double" w:sz="18" w:space="0" w:color="1E3A8A"/>'
        f'<w:left w:val="double" w:sz="18" w:space="0" w:color="1E3A8A"/>'
        f'<w:bottom w:val="double" w:sz="18" w:space="0" w:color="1E3A8A"/>'
        f'<w:right w:val="double" w:sz="18" w:space="0" w:color="1E3A8A"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)

    # 1. Institutional Header inside certificate box
    p_inst = outer_cell.paragraphs[0]
    p_inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_inst.paragraph_format.space_after = Pt(2)
    
    r_rep = p_inst.add_run("Republic of the Philippines\n")
    r_rep.font.size = Pt(9.5)
    r_rep.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)

    r_uni = p_inst.add_run("BUKIDNON STATE UNIVERSITY\n")
    r_uni.bold = True
    r_uni.font.size = Pt(15)
    r_uni.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)  # Navy

    r_col = p_inst.add_run("COLLEGE OF TECHNOLOGIES\n")
    r_col.bold = True
    r_col.font.size = Pt(11.5)
    r_col.font.color.rgb = RGBColor(0x92, 0x40, 0x0E)  # Gold/Amber Accent

    r_dep = p_inst.add_run("Department of Information Technology • Malaybalay City, Bukidnon\n")
    r_dep.font.size = Pt(9.5)
    r_dep.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    # Ornamental Divider
    p_div = outer_cell.add_paragraph()
    p_div.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_div.paragraph_format.space_before = Pt(4)
    p_div.paragraph_format.space_after = Pt(12)
    r_div = p_div.add_run("❖  ― ― ― ― ― ― ― ― ― ― ― ― ― ― ― ― ― ― ― ― ― ― ― ― ― ― ― ― ―  ❖")
    r_div.font.color.rgb = RGBColor(0xD9, 0x77, 0x06)  # Amber
    r_div.font.size = Pt(9)

    # 2. Certificate Title Banner
    p_cert = outer_cell.add_paragraph()
    p_cert.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cert.paragraph_format.space_after = Pt(12)
    r_cert = p_cert.add_run("CERTIFICATE OF CAPSTONE COMPLETION\n& FINAL RE-DEFENSE APPROVAL")
    r_cert.bold = True
    r_cert.font.size = Pt(14.5)
    r_cert.font.color.rgb = RGBColor(0x11, 0x18, 0x27)

    # 3. Certificate Body Text
    p_body = outer_cell.add_paragraph()
    p_body.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_body.paragraph_format.space_after = Pt(10)
    p_body.paragraph_format.line_spacing = 1.25
    p_body.add_run("This is to certify that the capstone research project entitled:\n\n")

    r_proj = p_body.add_run("“{{PROJECT_TITLE}}”\n\n")
    r_proj.bold = True
    r_proj.font.size = Pt(12.5)
    r_proj.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

    p_body.add_run(
        "prepared and successfully defended by the student researchers:\n\n"
    )

    r_authors = p_body.add_run("{{STUDENT_AUTHORS}}\n\n")
    r_authors.bold = True
    r_authors.font.size = Pt(11)
    r_authors.font.color.rgb = RGBColor(0x11, 0x18, 0x27)

    p_body.add_run(
        "has been comprehensively evaluated and recommended for oral defense approval by the "
        "Capstone Evaluation Committee, having satisfied all institutional requirements, "
        "rigorous dual-layer plagiarism verification (< 20% threshold), and complete multi-signatory "
        "Action Done Matrix (ADM) compliance for the degree of\n\n"
    )

    r_degree = p_body.add_run("BACHELOR OF SCIENCE IN INFORMATION TECHNOLOGY\n")
    r_degree.bold = True
    r_degree.font.size = Pt(11.5)
    r_degree.font.color.rgb = RGBColor(0x92, 0x40, 0x0E)

    r_date = p_body.add_run("Conferred and Approved this {{DATE_OF_APPROVAL}} at Bukidnon State University.\n")
    r_date.font.size = Pt(10)

    # 4. Signatories Grid Table
    p_sig_h = outer_cell.add_paragraph()
    p_sig_h.paragraph_format.space_before = Pt(8)
    p_sig_h.paragraph_format.space_after = Pt(6)
    p_sig_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sh = p_sig_h.add_run("CAPSTONE EVALUATION AND APPROVAL COMMITTEE")
    r_sh.bold = True
    r_sh.font.size = Pt(10)
    r_sh.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

    sig_table = outer_cell.add_table(rows=3, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sig_table.autofit = False

    s_col_w = [Inches(3.3), Inches(3.3)]
    signatories = [
        ("{{PANEL_CHAIR_NAME}}", "Research & Ethics Committee (REC) Chair\nDate: {{PANEL_CHAIR_SIGN_DATE}}"),
        ("{{ADVISER_NAME}}", "Capstone Project Adviser\nDate: {{ADVISER_SIGN_DATE}}"),
        ("{{PANEL_MEMBER_1_NAME}}", "Panel Member\nDate: {{MEMBER_1_SIGN_DATE}}"),
        ("{{PANEL_MEMBER_2_NAME}}", "Panel Member\nDate: {{MEMBER_2_SIGN_DATE}}"),
        ("{{DEFENSE_SECRETARY_NAME}}", "Defense Secretary / Document Custodian\nDate: {{SECRETARY_SIGN_DATE}}"),
        ("{{DEAN_NAME}}", "Dean, College of Technologies\nDate: {{DEAN_SIGN_DATE}}")
    ]

    for idx, (sig_name, sig_title) in enumerate(signatories):
        r_idx = idx // 2
        c_idx = idx % 2
        cell = sig_table.rows[r_idx].cells[c_idx]
        cell.width = s_col_w[c_idx]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(4)
        p.add_run("_____________________________________\n")
        r_name = p.add_run(sig_name + "\n")
        r_name.bold = True
        r_name.font.size = Pt(9.5)
        r_t = p.add_run(sig_title)
        r_t.font.size = Pt(8)
        r_t.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)

    # 5. Security Hash & Verification Seal
    p_sec = outer_cell.add_paragraph()
    p_sec.paragraph_format.space_before = Pt(12)
    p_sec.paragraph_format.space_after = Pt(2)
    p_sec.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_hash = p_sec.add_run(
        "Institutional Verification Hash: {{VERIFICATION_SHA256_HASH}}\n"
        "BukSU CMS-V2 Automated Certification Subsystem • Non-Repudiation Verified"
    )
    r_hash.font.size = Pt(7.5)
    r_hash.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    doc.save(OUTPUT_DOCX)
    print(f"[SUCCESS] Capstone Completion Certificate docx template successfully generated: {OUTPUT_DOCX}")


if __name__ == "__main__":
    build_completion_certificate()
