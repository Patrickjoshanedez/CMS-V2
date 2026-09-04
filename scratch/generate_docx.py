import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

MD_PATH = r"C:\Users\patri\.gemini\antigravity-ide\brain\a858ab0a-cdaa-4287-8955-de6cf0810704\architecture_audit.md"
DIAG1_PATH = r"c:\Users\patri\OneDrive\Desktop\Holy folder\CMS-V2\scratch\diagram1_topology.png"
DIAG2_PATH = r"c:\Users\patri\OneDrive\Desktop\Holy folder\CMS-V2\scratch\diagram2_plagiarism.png"
DIAG3_PATH = r"c:\Users\patri\OneDrive\Desktop\Holy folder\CMS-V2\scratch\diagram3_lifecycle.png"

OUTPUT_DOCX_DOCS = r"c:\Users\patri\OneDrive\Desktop\Holy folder\CMS-V2\docs\architecture\CMS_V2_Full_Architecture_Audit.docx"
OUTPUT_DOCX_BRAIN = r"C:\Users\patri\.gemini\antigravity-ide\brain\a858ab0a-cdaa-4287-8955-de6cf0810704\CMS_V2_Full_Architecture_Audit.docx"

# Color Palette Constants
COLOR_PRIMARY_NAVY = RGBColor(15, 23, 42)      # #0f172a
COLOR_ACCENT_BLUE = RGBColor(2, 132, 199)      # #0284c7
COLOR_TEXT_CHARCOAL = RGBColor(51, 65, 85)     # #334155
COLOR_MUTED = RGBColor(100, 116, 139)          # #64748b
HEX_HEADER_BG = "1E293B"                       # Deep slate navy
HEX_ZEBRA_BG = "F8FAFC"                        # Very light slate
HEX_CODE_BG = "F1F5F9"                         # Light gray for code
HEX_BORDER = "CBD5E1"                          # Subtle border

def set_cell_background(cell, hex_color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def set_cell_borders(cell, color="CBD5E1", sz="4", val="single"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        </w:tcBorders>
    ''')
    tcPr.append(tcBorders)

def add_styled_paragraph(doc, text="", style='Normal', space_after=6, space_before=0, line_spacing=1.15):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = line_spacing
    if text:
        append_formatted_text(p, text)
    return p

def append_formatted_text(paragraph, text, base_color=COLOR_TEXT_CHARCOAL, base_size=10.5, is_bold_all=False, is_italic_all=False):
    # Regex for **bold** and `code`
    tokens = re.split(r'(\*\*.*?\*\*|`.*?`|\*.*?\*)', text)
    for token in tokens:
        if not token:
            continue
        run = paragraph.add_run()
        run.font.name = 'Calibri'
        run.font.size = Pt(base_size)
        run.font.color.rgb = base_color

        if token.startswith('**') and token.endswith('**'):
            run.text = token[2:-2]
            run.bold = True
        elif token.startswith('*') and token.endswith('*'):
            run.text = token[1:-1]
            run.italic = True
        elif token.startswith('`') and token.endswith('`'):
            run.text = token[1:-1]
            run.font.name = 'Consolas'
            run.font.size = Pt(base_size - 1)
            run.font.color.rgb = RGBColor(14, 116, 144) # teal/slate
        else:
            run.text = token
            run.bold = is_bold_all
            run.italic = is_italic_all

def add_callout_box(doc, code_lines):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, HEX_CODE_BG)
    set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
    
    # Border: left thick accent, others thin
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="none"/>
            <w:left w:val="single" w:sz="24" w:space="0" w:color="0284C7"/>
            <w:bottom w:val="none"/>
            <w:right w:val="none"/>
        </w:tcBorders>
    ''')
    tcPr.append(tcBorders)

    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.05

    for i, line in enumerate(code_lines):
        if i > 0:
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(9.0)
        run.font.color.rgb = RGBColor(30, 41, 59)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def add_table_data(doc, headers, rows):
    col_count = len(headers)
    table = doc.add_table(rows=len(rows) + 1, cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Format Header
    for c_idx, h_text in enumerate(headers):
        cell = table.cell(0, c_idx)
        set_cell_background(cell, HEX_HEADER_BG)
        set_cell_margins(cell, top=140, bottom=140, left=140, right=140)
        set_cell_borders(cell, color="475569", sz="4")
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(h_text)
        run.font.name = 'Calibri'
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.bold = True

    # Format Data Rows
    for r_idx, row in enumerate(rows):
        is_zebra = (r_idx % 2 == 1)
        bg_color = HEX_ZEBRA_BG if is_zebra else "FFFFFF"
        for c_idx, val in enumerate(row):
            if c_idx >= col_count:
                continue
            cell = table.cell(r_idx + 1, c_idx)
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
            set_cell_borders(cell, color=HEX_BORDER, sz="4")
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.1
            append_formatted_text(p, val, base_size=9.0)

    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_after = Pt(6)

def build_docx():
    print("Reading markdown:", MD_PATH)
    with open(MD_PATH, "r", encoding="utf-8") as f:
        content = f.read()

    doc = Document()

    # Set document margins (1 inch)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Document Header / Cover Banner
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    run_t = title_p.add_run("🏛️ Full System Architecture Audit")
    run_t.font.name = 'Calibri'
    run_t.font.size = Pt(24)
    run_t.font.color.rgb = COLOR_PRIMARY_NAVY
    run_t.bold = True

    subtitle_p = doc.add_paragraph()
    subtitle_p.paragraph_format.space_before = Pt(0)
    subtitle_p.paragraph_format.space_after = Pt(12)
    run_st = subtitle_p.add_run("BukSU Capstone Management System V2 (CMS-V2)")
    run_st.font.name = 'Calibri'
    run_st.font.size = Pt(15)
    run_st.font.color.rgb = COLOR_ACCENT_BLUE
    run_st.bold = True

    meta_table = doc.add_table(rows=2, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_table.autofit = False

    meta_items = [
        ("Target Platform:", "BukSU Capstone Management System V2"),
        ("Compliance Standard:", "ASDLC v2.0-Ready (Agentic Lifecycle)"),
        ("Architecture Pattern:", "Monorepo (Express 5 + React 18 + FastAPI Microservice)"),
        ("Audit Status:", "Verified: 60/60 Agentic Checks | 100% Endpoint Parity")
    ]
    for idx, (label, val) in enumerate(meta_items):
        r = idx // 2
        c = idx % 2
        cell = meta_table.cell(r, c)
        cell.width = Inches(3.25)
        set_cell_background(cell, "F8FAFC")
        set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
        set_cell_borders(cell, color="E2E8F0", sz="4")
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        r1 = p.add_run(label + " ")
        r1.font.name = 'Calibri'
        r1.font.size = Pt(9.0)
        r1.font.color.rgb = COLOR_MUTED
        r1.bold = True
        r2 = p.add_run(val)
        r2.font.name = 'Calibri'
        r2.font.size = Pt(9.0)
        r2.font.color.rgb = COLOR_PRIMARY_NAVY

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Process markdown lines
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_lines = []
    in_mermaid_block = False
    mermaid_counter = 0

    in_table = False
    table_headers = []
    table_rows = []

    while i < len(lines):
        line = lines[i]

        # Check for mermaid block
        if line.strip().startswith('```mermaid'):
            in_mermaid_block = True
            mermaid_counter += 1
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                i += 1
            # After mermaid block ends, insert corresponding diagram image!
            diag_path = None
            diag_caption = ""
            if mermaid_counter == 1:
                diag_path = DIAG1_PATH
                diag_caption = "Figure 1: High-Level System Architecture & Component Topology"
            elif mermaid_counter == 2:
                diag_path = DIAG2_PATH
                diag_caption = "Figure 2: Dual-Engine Plagiarism & Originality Detection Architecture"
            elif mermaid_counter == 3:
                diag_path = DIAG3_PATH
                diag_caption = "Figure 3: Canonical 4-Phase Institutional Capstone Lifecycle Progression (Phase 0 to Phase 4)"

            if diag_path and os.path.exists(diag_path):
                img_p = doc.add_paragraph()
                img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                img_p.paragraph_format.space_before = Pt(10)
                img_p.paragraph_format.space_after = Pt(4)
                run_img = img_p.add_run()
                run_img.add_picture(diag_path, width=Inches(6.2))

                cap_p = doc.add_paragraph()
                cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap_p.paragraph_format.space_before = Pt(2)
                cap_p.paragraph_format.space_after = Pt(12)
                run_cap = cap_p.add_run(diag_caption)
                run_cap.font.name = 'Calibri'
                run_cap.font.size = Pt(9.0)
                run_cap.font.color.rgb = COLOR_MUTED
                run_cap.italic = True
            
            in_mermaid_block = False
            i += 1
            continue

        # Check for standard code block
        if line.strip().startswith('```') and not in_mermaid_block:
            if not in_code_block:
                in_code_block = True
                code_lines = []
                i += 1
                continue
            else:
                in_code_block = False
                if code_lines:
                    add_callout_box(doc, code_lines)
                code_lines = []
                i += 1
                continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Check for Markdown Tables
        if '|' in line and not in_code_block:
            raw_cells = [c.strip() for c in line.split('|')[1:-1]]
            if raw_cells:
                # Check if it's separator row: | :--- | :--- |
                if all(re.match(r'^:?-+:?$', c) for c in raw_cells):
                    i += 1
                    continue
                if not in_table:
                    in_table = True
                    table_headers = raw_cells
                    table_rows = []
                else:
                    table_rows.append(raw_cells)
                i += 1
                continue
        else:
            if in_table:
                add_table_data(doc, table_headers, table_rows)
                in_table = False
                table_headers = []
                table_rows = []

        # Headings
        if line.startswith('# '):
            # Already added cover title, skip title or format as H1
            if "🏛️ Full System Architecture Audit" in line:
                i += 1
                continue
            h = doc.add_heading(level=1)
            h.paragraph_format.space_before = Pt(14)
            h.paragraph_format.space_after = Pt(6)
            run = h.add_run(line[2:].strip())
            run.font.name = 'Calibri'
            run.font.size = Pt(16)
            run.font.color.rgb = COLOR_PRIMARY_NAVY
            run.bold = True
            i += 1
            continue
        elif line.startswith('## '):
            h = doc.add_heading(level=2)
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(5)
            run = h.add_run(line[3:].strip())
            run.font.name = 'Calibri'
            run.font.size = Pt(13.5)
            run.font.color.rgb = COLOR_PRIMARY_NAVY
            run.bold = True
            i += 1
            continue
        elif line.startswith('### '):
            h = doc.add_heading(level=3)
            h.paragraph_format.space_before = Pt(10)
            h.paragraph_format.space_after = Pt(4)
            run = h.add_run(line[4:].strip())
            run.font.name = 'Calibri'
            run.font.size = Pt(11.5)
            run.font.color.rgb = COLOR_ACCENT_BLUE
            run.bold = True
            i += 1
            continue
        elif line.startswith('#### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(line[5:].strip())
            run.font.name = 'Calibri'
            run.font.size = Pt(10.5)
            run.font.color.rgb = COLOR_PRIMARY_NAVY
            run.bold = True
            i += 1
            continue

        # Bullet Lists
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            clean_text = line.strip()[2:]
            append_formatted_text(p, clean_text)
            i += 1
            continue

        # Horizontal rule
        if line.strip() in ['---', '***', '___']:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            pBrd = parse_xml(f'<w:pBrd {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="E2E8F0"/></w:pBrd>')
            p._p.get_or_add_pPr().append(pBrd)
            i += 1
            continue

        # Normal Paragraph
        if line.strip():
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.15
            append_formatted_text(p, line.strip())

        i += 1

    # End of document check for remaining table
    if in_table:
        add_table_data(doc, table_headers, table_rows)

    os.makedirs(os.path.dirname(OUTPUT_DOCX_DOCS), exist_ok=True)
    saved_docs_path = OUTPUT_DOCX_DOCS
    try:
        doc.save(OUTPUT_DOCX_DOCS)
        print("Saved DOCX to docs:", OUTPUT_DOCX_DOCS)
    except PermissionError:
        fallback_path = os.path.join(os.path.dirname(OUTPUT_DOCX_DOCS), "CMS_V2_Full_Architecture_Audit_Updated.docx")
        doc.save(fallback_path)
        saved_docs_path = fallback_path
        print(f"Original file was locked by an application (e.g. Word). Saved DOCX to fallback: {fallback_path}")

    os.makedirs(os.path.dirname(OUTPUT_DOCX_BRAIN), exist_ok=True)
    try:
        doc.save(OUTPUT_DOCX_BRAIN)
        print("Saved DOCX to brain artifacts:", OUTPUT_DOCX_BRAIN)
    except PermissionError:
        fallback_brain = os.path.join(os.path.dirname(OUTPUT_DOCX_BRAIN), "CMS_V2_Full_Architecture_Audit_Updated.docx")
        doc.save(fallback_brain)
        print(f"Saved DOCX to brain fallback: {fallback_brain}")

if __name__ == "__main__":
    build_docx()
